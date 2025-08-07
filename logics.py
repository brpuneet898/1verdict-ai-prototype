import os
import json
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.docstore.document import Document
from langchain.chains.question_answering import load_qa_chain
from langchain.chains import LLMChain
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import JsonOutputParser
import PyPDF2
import docx
import yaml
# --- Constants ---
EMBEDDING_MODEL = "sentence-transformers/all-MiniLM-L6-v2"

def extract_text_from_file(filepath):
    """Extracts text from PDF or DOCX files."""
    ext = filepath.rsplit('.', 1)[1].lower()
    text = ""
    if ext == "pdf":
        try:
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() or ""
        except Exception as e:
            print(f"Error reading PDF {filepath}: {e}")
            raise
    elif ext == "docx":
        try:
            doc = docx.Document(filepath)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            print(f"Error reading DOCX {filepath}: {e}")
            raise
    else:
        raise ValueError("Unsupported file type")
    return text

def get_text_chunks(text):
    """Splits text into manageable chunks."""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )
    chunks = text_splitter.split_text(text)
    # Convert chunks to Document objects for LangChain
    documents = [Document(page_content=chunk) for chunk in chunks]
    return documents

def get_vector_store(documents):
    """Creates a FAISS vector store from text chunks."""
    try:
        embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)
        vector_store = FAISS.from_documents(documents, embedding=embeddings)
        return vector_store
    except Exception as e:
        print(f"Error creating vector store: {e}")
        raise

# This is the corrected version of your original function
def get_summary_from_llm(llm, vector_store):
    """Queries the LLM to get a summary."""
    
    try:
        with open("key.yaml", "r") as f:
            config = yaml.safe_load(f)
        api_key = config.get("GEMINI_API_KEY")
        if not api_key:
            raise ValueError("GEMINI_API_KEY not found in key.yaml.")
    except FileNotFoundError:
        raise FileNotFoundError("key.yaml not found. Please ensure it is in the root directory.")
    except Exception as e:
        raise e

    model_name = llm.model_name.split('/')[-1]
    chat_llm_wrapper = ChatGoogleGenerativeAI(
        model=model_name,
        google_api_key=api_key,
        temperature=0.3,
        convert_system_message_to_human=True
    )


    prompt_template_str = """
    Based on the provided document, please generate a concise summary.
    The summary should be between 150 and 200 words.
    Focus on the main points, key arguments, and conclusions.
    
    Document context is provided below.

    CONTEXT:
    {context}

    SUMMARY:
    """
    prompt = PromptTemplate.from_template(prompt_template_str)
    
    chain = load_qa_chain(chat_llm_wrapper, chain_type="stuff", prompt=prompt)
    
    query = "Summarize the entire document."
    docs = vector_store.similarity_search(query, k=5)

    if not docs:
        return "Could not find any relevant text to summarize."

    response = chain.invoke({"input_documents": docs, "question": query})
    
    return response.get('output_text', 'Failed to generate summary.')


def summarize_text(filepath, llm):
    """
    Main function to orchestrate the summarization process.
    """
    print(f"Starting summarization for: {filepath}")

    raw_text = extract_text_from_file(filepath)
    if not raw_text.strip():
        return "Could not extract text from the document. It might be empty or scanned."

    documents = get_text_chunks(raw_text)
    if not documents:
        return "Failed to create text chunks from the document."

    vector_store = get_vector_store(documents)

    summary = get_summary_from_llm(llm, vector_store)

    return summary


def get_clauses_from_chain(llm, vector_store):
    """Identifies and extracts all clauses from the document using a direct LLMChain."""
    try:
        with open("key.yaml", "r") as f:
            config = yaml.safe_load(f)
        api_key = config.get("GEMINI_API_KEY")
        if not api_key: raise ValueError("GEMINI_API_KEY not found in key.yaml.")
    except FileNotFoundError:
        raise FileNotFoundError("key.yaml not found.")

    model_name = llm.model_name.split('/')[-1]
    chat_llm_wrapper = ChatGoogleGenerativeAI(model=model_name, google_api_key=api_key, temperature=0.3, convert_system_message_to_human=True)
    
    # --- THIS IS THE FIX ---
    # We create a simpler prompt that only needs the document context.
    # This avoids the complex input key mapping that caused the error.
    output_parser = JsonOutputParser()

    prompt_template_str = """
    Analyze the following document text and identify all distinct legal or policy clauses.
    For each clause you find, provide a title and the full, extracted text of that clause.
    Your output MUST be a valid JSON array, where each object has a "title" and a "text" key.
    
    {format_instructions}

    DOCUMENT TEXT:
    {context}
    
    JSON ARRAY:
    """
    prompt = PromptTemplate(
        template=prompt_template_str,
        input_variables=["context"],
        partial_variables={"format_instructions": output_parser.get_format_instructions()}
    )

    # This is the modern LangChain Expression Language (LCEL) syntax.
    # We "pipe" the components together instead of using the deprecated LLMChain.
    chain = prompt | chat_llm_wrapper | output_parser
    
    query = "Extract all legal and policy clauses from the document."
    docs = vector_store.similarity_search(query, k=10)
    if not docs: return []

    context_string = "\n\n".join([doc.page_content for doc in docs])

    try:
        # Run the chain with the single 'context' input. The parser will handle the output.
        clauses = chain.invoke({"context": context_string})
        return clauses
    except Exception as e:
        # This will catch errors if the model output is not valid JSON.
        print(f"Error parsing JSON from model output: {e}")
        return [{"title": "Parsing Error", "text": "The AI returned a response that could not be read as valid JSON."}]

def review_key_clauses(filepath, llm):
    """Main function to orchestrate the clause review process."""
    print(f"Starting clause review for: {filepath}")
    raw_text = extract_text_from_file(filepath)
    if not raw_text.strip(): return []
    
    text_chunks = get_text_chunks(raw_text)
    if not text_chunks: return []
        
    vector_store = get_vector_store(text_chunks)
    clauses = get_clauses_from_chain(llm, vector_store)
    return clauses
